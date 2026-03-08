"""Write a structured Word document with optional references section and proper citation formatting."""

from pathlib import Path

from hivemind.tools.base import Tool
from hivemind.tools.registry import register


def _add_hyperlink(paragraph, text: str, url: str):
    """Add a hyperlink run to a paragraph (python-docx has no built-in add_hyperlink)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(u)
    r_pr.append(color)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_reference_paragraph(doc, number: int, ref: dict) -> None:
    """Add one reference as a numbered paragraph with APA-like formatting and optional hyperlink."""
    from docx.shared import Pt

    authors = ref.get("authors") or ref.get("author") or "Unknown"
    title = ref.get("title") or ""
    year = ref.get("year") or ""
    journal = ref.get("journal")
    url = ref.get("url")
    publisher = ref.get("publisher")

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.first_line_indent = Pt(-18)
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run(f"[{number}] ")
    run.bold = True

    para.add_run(f"{authors} ({year}). ")

    r_title = para.add_run(f"{title}.")
    r_title.italic = True
    para.add_run(" ")

    if journal:
        r_journal = para.add_run(f" {journal}.")
        r_journal.italic = True
    if publisher:
        para.add_run(f" {publisher}.")
    if url:
        para.add_run(" ")
        _add_hyperlink(para, url if len(url) < 50 else "Link", url)


class WriteWordDocumentTool(Tool):
    """Write a structured Word document with optional references. Use [1], [2] in content to match numbered refs."""

    name = "write_word_document"
    description = (
        "Write a structured Word document with optional references. "
        "Use [1] / Author (Year) (Word/MD) in content; pass references array (key, authors, title, year, journal, url, publisher) for bibliography. Produces .docx. Requires: pip install hivemind-ai[document]."
    )
    input_schema = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Output path (e.g. report.docx)"},
            "title": {"type": "string", "description": "Document title"},
            "content": {"type": "string", "description": "Body (plain or markdown-style newlines). Use [1], [2] for citations."},
            "references": {
                "type": "array",
                "items": {"type": "object"},
                "description": "Optional list of refs: key, authors, title, year, journal, url, publisher",
            },
        },
        "required": ["path", "title", "content"],
    }

    def run(self, **kwargs) -> str:
        try:
            from docx import Document
        except ImportError:
            return "Word export requires the document extra: pip install hivemind-ai[document]"

        path = kwargs.get("path")
        title = kwargs.get("title")
        content = kwargs.get("content")
        references = kwargs.get("references")

        if not path or not isinstance(path, str) or not path.strip():
            return "Error: path must be a non-empty string"
        if not title or not isinstance(title, str):
            return "Error: title must be a non-empty string"
        if content is None:
            return "Error: content is required"
        if not isinstance(content, str):
            content = str(content)

        path = path.strip()
        if not path.endswith(".docx"):
            path = path.rstrip("/") + ".docx"
        p = Path(path).resolve()
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            return f"Error creating directory: {e}"

        doc = Document()
        doc.add_heading(title.strip(), level=0)
        for block in content.strip().split("\n\n"):
            block = block.strip()
            if block:
                doc.add_paragraph(block)

        if references and isinstance(references, list):
            doc.add_heading("References", level=1)
            num = 0
            for ref in references:
                if not isinstance(ref, dict):
                    continue
                num += 1
                _add_reference_paragraph(doc, num, ref)

        try:
            doc.save(str(p))
        except Exception as e:
            return f"Error writing file: {e}"
        return f"Wrote {p}"


register(WriteWordDocumentTool())
